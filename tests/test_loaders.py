"""
Unit tests for KGClaw document loaders.

Tests CSV (co-occurrence network data), JSONL (character relationships),
TXT (character network), XLSX, and PDF loaders with diverse test data.
"""

import csv
import io
import json
import shutil
import tempfile
from pathlib import Path

import pytest

from kgclaw.loaders import (
    LoadedDocument,
    get_loader,
    load_csv,
    load_docx,
    load_html,
    load_jsonl,
    load_pdf,
    load_text,
    load_xlsx,
    load_directory,
    supported_extensions,
)


class TestLoaders:
    """Tests for all document loaders with diverse test data."""

    @pytest.fixture(autouse=True)
    def setup(self):
        self.tmp = Path(tempfile.mkdtemp())
        yield
        shutil.rmtree(self.tmp, ignore_errors=True)

    # ── TXT Loader ───────────────────────────────────────────────────────

    def test_load_text_basic(self):
        f = self.tmp / "test.txt"
        f.write_text("赵铁蛋是赵本山的儿子。\n赵玉田是赵本山的儿子。")
        doc = load_text(str(f))
        assert "赵铁蛋" in doc.text
        assert doc.source == str(f)
        assert doc.metadata["ext"] == ".txt"
        assert doc.metadata["size"] > 0

    def test_load_text_markdown(self):
        f = self.tmp / "README.md"
        f.write_text("# 人物关系\n\n- 赵铁蛋\n- 赵本山")
        doc = load_text(str(f))
        assert "人物关系" in doc.text
        assert doc.metadata["ext"] == ".md"

    def test_load_text_empty_file(self):
        f = self.tmp / "empty.txt"
        f.write_text("")
        doc = load_text(str(f))
        assert doc.text == ""

    # ── CSV Loader (co-occurrence network data) ──────────────────────────

    def test_load_csv_co_occurrence_network(self):
        """Test CSV with co-occurrence network data."""
        csv_content = """source,target,weight,context
赵铁蛋,赵本山,5,父子关系
赵玉田,赵本山,3,父子关系
刘海柱,赵铁蛋,2,朋友关系
赵铁蛋,李大明,1,同事
赵玉田,刘海柱,4,商业合作
"""
        f = self.tmp / "co_occurrence_network.csv"
        f.write_text(csv_content)

        doc = load_csv(str(f))
        assert doc.source == str(f)
        assert doc.metadata["is_tabular"] is True
        assert doc.metadata["rows"] == 5
        assert doc.metadata["columns"] == ["source", "target", "weight", "context"]

        # Check raw_rows preserved for structured extraction
        raw_rows = doc.metadata["raw_rows"]
        assert len(raw_rows) == 5
        assert raw_rows[0]["source"] == "赵铁蛋"
        assert raw_rows[0]["weight"] == "5"

        # Check text representation
        assert "赵铁蛋" in doc.text
        assert "赵本山" in doc.text
        assert "Columns:" in doc.text

    def test_load_csv_with_bom(self):
        """CSV with UTF-8 BOM."""
        f = self.tmp / "bom_test.csv"
        # Write with BOM
        with open(f, "w", encoding="utf-8-sig") as fh:
            fh.write("姓名,年龄,城市\n张三,30,北京\n李四,25,上海")
        doc = load_csv(str(f))
        assert doc.metadata["rows"] == 2
        assert doc.metadata["columns"] == ["姓名", "年龄", "城市"]

    def test_load_csv_tsv(self):
        """TSV file with tab delimiter."""
        f = self.tmp / "data.tsv"
        f.write_text("name\tvalue\tcategory\nAlice\t100\tA\nBob\t200\tB")
        doc = load_csv(str(f))
        assert doc.metadata["rows"] == 2
        assert doc.metadata["columns"] == ["name", "value", "category"]

    def test_load_csv_large_dataset_capped_at_500(self):
        """Large CSV — raw_rows preserved, text capped at 500."""
        f = self.tmp / "large.csv"
        with open(f, "w") as fh:
            fh.write("id,name,value\n")
            for i in range(600):
                fh.write(f"{i},item_{i},{i * 10}\n")
        doc = load_csv(str(f))
        assert doc.metadata["rows"] == 600
        assert len(doc.metadata["raw_rows"]) == 600
        # Text representation should mention truncation
        assert "500" in doc.text or "more rows" in doc.text

    def test_load_csv_empty(self):
        """Empty CSV (only headers)."""
        f = self.tmp / "empty.csv"
        f.write_text("col1,col2")
        doc = load_csv(str(f))
        assert doc.metadata["rows"] == 0

    def test_load_csv_malformed_fallback_to_text(self):
        """Corrupted CSV falls back to plain text."""
        f = self.tmp / "bad.csv"
        # Write binary content that confuses DictReader
        f.write_bytes(b"\x00\x01\x02\x03\x04")
        doc = load_csv(str(f))
        assert doc.source == str(f)
        # Should still produce something
        assert doc.text is not None

    # ── JSONL Loader (character relationships) ───────────────────────────

    def test_load_jsonl_character_relationships(self):
        """JSONL with character relationship extraction data."""
        lines = [
            {"text": "赵铁蛋是赵本山的儿子", "entities": ["赵铁蛋", "赵本山"], "relation": "生父"},
            {"text": "赵玉田住在象牙山村", "entities": ["赵玉田", "象牙山村"], "relation": "居住于"},
            {"text": "刘海柱与赵铁蛋是朋友", "entities": ["刘海柱", "赵铁蛋"], "relation": "朋友"},
        ]
        f = self.tmp / "character_relations.jsonl"
        content = "\n".join(json.dumps(line, ensure_ascii=False) for line in lines)
        f.write_text(content)

        doc = load_jsonl(str(f))
        assert doc.source == str(f)
        assert doc.metadata["lines"] == 3
        assert "赵铁蛋是赵本山的儿子" in doc.text
        assert "赵玉田" in doc.text
        assert "刘海柱" in doc.text

    def test_load_jsonl_empty_lines(self):
        """JSONL with empty lines between entries."""
        content = '{"text": "hello"}\n\n{"text": "world"}\n'
        f = self.tmp / "sparse.jsonl"
        f.write_text(content)
        doc = load_jsonl(str(f))
        assert doc.metadata["lines"] == 2

    def test_load_jsonl_invalid_lines(self):
        """JSONL with some invalid lines."""
        content = '{"text": "hello"}\nnot valid json\n{"text": "world"}'
        f = self.tmp / "mixed.jsonl"
        f.write_text(content)
        doc = load_jsonl(str(f))
        # Loader appends raw line on JSON decode error, so all 3 lines are counted
        assert doc.metadata["lines"] == 3
        assert "hello" in doc.text
        assert "world" in doc.text
        assert "not valid json" in doc.text

    def test_load_jsonl_multiple_fields(self):
        """JSONL with varied field structure — only 'text' or 'data' field extracted."""
        entries = [
            {"text": "doc1", "source": "src1", "date": "2024-01-01", "tags": ["a", "b"]},
            {"text": "doc2", "source": "src2", "author": "张三", "score": 0.95},
        ]
        f = self.tmp / "rich.jsonl"
        f.write_text("\n".join(json.dumps(e, ensure_ascii=False) for e in entries))
        doc = load_jsonl(str(f))
        # Loader extracts only the "text" field (or "data" as fallback)
        assert doc.metadata["lines"] == 2
        assert "doc1" in doc.text
        assert "doc2" in doc.text
        # Other fields (source, date, author) are NOT extracted into text
        # because load_jsonl only looks for "data" or "text" keys

    # ── XLSX Loader ──────────────────────────────────────────────────────

    def test_load_xlsx_basic(self):
        """Test XLSX loading if openpyxl is available."""
        try:
            from openpyxl import Workbook
        except ImportError:
            pytest.skip("openpyxl not installed")

        f = self.tmp / "test.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.title = "人物关系"
        ws.append(["姓名", "关系", "对象"])
        ws.append(["赵铁蛋", "生父", "赵本山"])
        ws.append(["赵玉田", "居住于", "象牙山村"])
        wb.save(str(f))

        doc = load_xlsx(str(f))
        assert doc.metadata["is_tabular"] is True
        assert doc.metadata["rows"] == 2
        # raw_rows preserved
        assert len(doc.metadata["raw_rows"]) == 2
        assert doc.metadata["raw_rows"][0]["姓名"] == "赵铁蛋"

    def test_load_xlsx_multiple_sheets(self):
        """XLSX with multiple sheets."""
        try:
            from openpyxl import Workbook
        except ImportError:
            pytest.skip("openpyxl not installed")

        f = self.tmp / "multi_sheet.xlsx"
        wb = Workbook()
        ws1 = wb.active
        ws1.title = "实体"
        ws1.append(["id", "name"])
        ws1.append(["1", "赵铁蛋"])

        ws2 = wb.create_sheet("关系")
        ws2.append(["source", "target", "type"])
        ws2.append(["1", "2", "父子"])

        wb.save(str(f))

        doc = load_xlsx(str(f))
        assert doc.metadata["sheets"] == 2

    # ── DOCX Loader ──────────────────────────────────────────────────────

    def test_load_docx_basic(self):
        """Test DOCX loading if python-docx is available."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            pytest.skip("python-docx not installed")

        f = self.tmp / "test.docx"
        d = DocxDocument()
        d.add_paragraph("赵铁蛋是赵本山的儿子。")
        d.add_paragraph("赵玉田是赵铁蛋的兄弟。")
        d.save(str(f))

        doc = load_docx(str(f))
        assert "赵铁蛋" in doc.text
        assert "赵本山" in doc.text

    def test_load_docx_empty(self):
        """Empty DOCX."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            pytest.skip("python-docx not installed")

        f = self.tmp / "empty.docx"
        DocxDocument().save(str(f))

        doc = load_docx(str(f))
        assert doc.source == str(f)

    # ── PDF Loader ───────────────────────────────────────────────────────

    def test_load_pdf_basic(self):
        """Test PDF loading if PyPDF2 or pdfplumber is available."""
        # Try to create a minimal PDF
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import A4
        except ImportError:
            pytest.skip("reportlab not installed (needed to create test PDF)")

        f = self.tmp / "test.pdf"
        c = canvas.Canvas(str(f), pagesize=A4)
        c.drawString(100, 750, "人物关系图谱分析报告")
        c.drawString(100, 730, "赵铁蛋是赵本山的儿子")
        c.drawString(100, 710, "赵玉田居住在象牙山村")
        c.save()

        # Try loading
        try:
            doc = load_pdf(str(f))
            assert doc.source == str(f)
            assert doc.metadata["ext"] == ".pdf"
        except ImportError:
            pytest.skip("No PDF reader available (PyPDF2/pdfplumber)")

    def test_load_pdf_with_pypdf2(self):
        """PDF content extraction via PyPDF2."""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import A4
        except ImportError:
            pytest.skip("reportlab not installed")

        f = self.tmp / "char_report.pdf"
        c = canvas.Canvas(str(f), pagesize=A4)
        c.drawString(100, 750, "119届国会数据分析")
        c.drawString(100, 730, "众议员张三，来自加州第12选区")
        c.drawString(100, 710, "支持法案HR1234和HR5678")
        c.save()

        doc = load_pdf(str(f))
        # Even if no reader available, should not crash
        assert doc.source == str(f)

    # ── HTML Loader ──────────────────────────────────────────────────────

    def test_load_html_basic(self):
        """Test HTML loading if BeautifulSoup is available."""
        try:
            from bs4 import BeautifulSoup  # noqa: F401
        except ImportError:
            pytest.skip("beautifulsoup4 not installed")

        f = self.tmp / "page.html"
        f.write_text("""<!DOCTYPE html>
<html><body>
<h1>人物关系</h1>
<p>赵铁蛋是赵本山的儿子。</p>
<p>赵玉田是赵本山的儿子。</p>
<script>console.log('ignore this')</script>
<nav>skip nav</nav>
</body></html>""")

        doc = load_html(str(f))
        assert "赵铁蛋" in doc.text
        assert "赵本山" in doc.text
        # Script/nav content should be stripped
        assert "ignore this" not in doc.text.lower()
        assert "skip nav" not in doc.text.lower()

    # ── Directory Loader ─────────────────────────────────────────────────

    def test_load_directory_recursive(self):
        """Load a directory with mixed file types."""
        sub = self.tmp / "data"
        sub.mkdir()
        (sub / "a.txt").write_text("doc a content")
        (sub / "b.txt").write_text("doc b content")
        (sub / "notes.md").write_text("# Notes")

        docs = load_directory(str(sub), recursive=True)
        assert len(docs) >= 3
        texts = {d.text for d in docs}
        assert "doc a content" in texts
        assert "doc b content" in texts

    def test_load_directory_excludes_hidden(self):
        """Directory loader excludes .git, __pycache__, etc."""
        sub = self.tmp / "project"
        sub.mkdir()
        (sub / "data.txt").write_text("actual data")
        git_dir = sub / ".git"
        git_dir.mkdir()
        (git_dir / "config").write_text("git config")

        docs = load_directory(str(sub), recursive=True,
                              exclude_patterns=[".git/**"])
        texts = {d.text for d in docs}
        assert "actual data" in texts
        # .git/config should not appear
        assert not any(".git" in d.source for d in docs)

    # ── Loader Discovery ─────────────────────────────────────────────────

    def test_get_loader_for_all_supported_extensions(self):
        """Every supported extension has a loader."""
        for ext in supported_extensions():
            loader = get_loader(ext)
            assert loader is not None, f"No loader for {ext}"

    def test_get_loader_unknown_extension(self):
        """Unknown extension returns None."""
        assert get_loader(".xyz_unknown") is None

    def test_supported_extensions_includes_common(self):
        exts = supported_extensions()
        assert ".txt" in exts
        assert ".csv" in exts
        assert ".jsonl" in exts
